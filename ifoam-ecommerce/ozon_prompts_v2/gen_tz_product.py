#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gen_tz_product.py — генератор ТЗ на анализ и запуск карточки товара Ozon
Структура: Шаблон_ТЗ.docx (5 разделов + 2 таблицы)

Запуск из командной строки:
    python3 gen_tz_product.py --art 770758 --brand "ifoamHOME" \
        --name "КондиционерBalmy_РайскиеЦветы" \
        --product_name "balmy Кондиционер для белья «Райские цветы»" \
        --volume "1 л" --date "31.03.2026"

Или импорт функции:
    from gen_tz_product import build_product_tz
    build_product_tz(folder, art, brand, name, product_name, volume, date_str)
"""

import sys, argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8', errors='replace')


# ─── Хелперы ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _run(para, text, bold=False, size=10, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12 if level == 1 else 10)
    if level == 1:
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def kv(doc, key, value, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    _run(p, key + ': ', bold=True,  size=size)
    _run(p, value,       bold=False, size=size)


def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.add_run(text).font.size = Pt(size)


def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Шапка
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, 'D6E4F0')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    # Данные
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = t.rows[ri + 1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(str(val)).font.size = Pt(9)
    # Ширина колонок
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return t


# ─── Основная функция ────────────────────────────────────────────────────────

def build_product_tz(
    folder: Path,
    art: str,
    brand: str,
    name: str,           # slug для имени файла, напр. "КондиционерBalmy_РайскиеЦветы"
    product_name: str,   # человекочитаемое название, напр. 'balmy Кондиционер «Райские цветы»'
    volume: str,
    date_str: str,
    extra_chars: dict = None,   # дополнительные характеристики {ключ: значение}
) -> Path:
    """
    Генерирует 01a_tz_[art].docx в папке folder.
    Возвращает Path сохранённого файла.
    """
    folder = Path(folder)
    folder.mkdir(parents=True, exist_ok=True)
    out_path = folder / f'01a_tz_{art}.docx'

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21);    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5);   sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2);     sec.bottom_margin = Cm(2)

    # ── ЗАГОЛОВОК ──────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(3)
    r = p_title.add_run('ТЕХНИЧЕСКОЕ ЗАДАНИЕ')
    r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(3)
    p_sub.add_run(
        f'Анализ и запуск карточки товара {product_name} {volume} на Ozon (арт. {art})'
    ).font.size = Pt(11)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_after = Pt(14)
    r3 = p_date.add_run(f'Дата: {date_str}')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # ── 1. СТОРОНЫ ─────────────────────────────────────────────────────────
    heading(doc, '1. Стороны')
    kv(doc, 'Заказчик',             f'{brand} — менеджер по товарному контенту')
    kv(doc, 'Исполнитель',          'Аналитик / Claude Code (AI-агент) + оператор')
    kv(doc, 'Кто будет использовать',
       'Менеджер контента, дизайнер карточки, фотограф, SEO-специалист')

    # ── 2. ЦЕЛЬ ────────────────────────────────────────────────────────────
    heading(doc, '2. Цель задачи')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        f'Провести полный предстартовый анализ товара {product_name} {volume} '
        f'(арт. {art}, бренд {brand}) и сформировать комплект из 10 документов '
        f'для запуска конкурентоспособной карточки на Ozon.\n\n'
        f'Каждый шаг системы оформляется отдельным файлом (DOCX / XLSX / PPTX) '
        f'и сохраняется в папку продукта. Выход каждого шага — вход следующего.'
    ).font.size = Pt(10)

    # ── 3. ВВОДНЫЕ ДАННЫЕ ──────────────────────────────────────────────────
    heading(doc, '3. Вводные данные')

    base_folder = f'C:/Users/1/Downloads/{brand}_{name}_{art}/'
    vvod = [
        ('1', f'Этикетка товара → 01_card_{art}.docx (уже выполнено)',
         'Файл',    'Папка продукта (обработана Промптом 1)'),
        ('2', 'История продаж XLSX',
         'Таблица', 'DataLens или\nD:/Аналитика/Система/combined_marketplaces_with_profit.xlsx'),
        ('3', 'Топ конкурентов из Ozon Bestsellers XLSX',
         'Таблица', 'seller.ozon.ru → Аналитика → Конкуренты → Экспортировать'),
        ('4', 'Страницы отзывов топ-5 конкурентов (PDF × 5)',
         'Файл',    'Chrome → Ctrl+P → Сохранить как PDF'),
        ('5', 'Страница всех поисковых запросов (PDF)',
         'Файл',    'seller.ozon.ru → Аналитика → Поисковые запросы → Все → Ctrl+P'),
        ('6', 'Карточки конкурентов (PDF × 5)',
         'Файл',    'Chrome → Ctrl+P → Сохранить как PDF'),
        ('7', 'combined_reference_normalized.csv',
         'Система', 'D:/Аналитика/Система/\n(себестоимость, объём, категория)'),
        ('8', 'ozon_products_*.xlsx (свежайший)',
         'Система', 'D:/Аналитика/Система/Управленческий_отчет_IFOAM_ОЗОН/Товары/'),
    ]
    make_table(doc, ['#', 'Название', 'Тип', 'Источник'], vvod,
               col_widths=[0.7, 6.5, 2.0, 7.3])
    doc.add_paragraph()

    # ── 4. ЭТАПЫ ───────────────────────────────────────────────────────────
    heading(doc, '4. Этапы выполнения')
    etapy = [
        ('1',  f'Извлечение данных из этикетки (Промпт 1)\n✓ ВЫПОЛНЕНО',
               f'01_card_{art}.docx', 'Выполнено'),
        ('1a', 'ТЗ на анализ товара — настоящий документ (Промпт 1a)',
               f'01a_tz_{art}.docx', 'Выполнено'),
        ('1б', 'Инструкция по заполнению характеристик в Ozon Seller (Промпт 1b)',
               'Гайд для оператора (текст)', 'После Шага 1'),
        ('2',  'Анализ истории продаж (Промпт 2)',
               f'02_sales_history_{art}.docx', 'После получения XLSX'),
        ('3',  'Анализ конкурентов → таблица (Промпт 3)\nСтратегия ценообразования (Промпт 3b)',
               f'03_competitors_{art}.xlsx\n03b_pricing_strategy_{art}.docx',
               'После получения XLSX конкурентов'),
        ('4',  'Анализ отзывов конкурентов → боли и желания (Промпт 4)',
               f'04_pains_{art}.docx', 'После сохранения PDF отзывов'),
        ('5',  'SEO-анализ поисковых запросов (Промпт 5)',
               f'05_seo_{art}.xlsx + .docx', 'После сохранения PDF запросов'),
        ('6',  'Визуальный анализ карточек конкурентов (Промпт 6)',
               f'06_visual_{art}.docx', 'После сохранения PDF карточек'),
        ('7',  'ТЗ для дизайнера (Промпт 7)\nТЗ для фотографа (Промпт 7b)',
               f'07_designer_tz_{art}.docx\n07_designer_tz_{art}.pptx\nТЗ_Фотограф_{art}.docx',
               'После шагов 1–6'),
        ('8',  'Полный листинг товара (Промпт 8)\nИнструкция оператору (Промпт 8b)',
               f'08_listing_{art}.docx\nГайд для оператора', 'После всех шагов'),
    ]
    make_table(doc, ['№', 'Этап', 'Результат', 'Срок'], etapy,
               col_widths=[0.7, 6.2, 5.8, 3.8])
    doc.add_paragraph()

    # ── 5. УСЛОВИЯ ВЫПОЛНЕНИЯ ──────────────────────────────────────────────
    heading(doc, '5. Задача считается выполненной, когда:')
    conditions = [
        f'Все 10 выходных документов сохранены в папке {base_folder}',
        'Данные этикетки (дозировка, пропорции разбавления, состав, pH, характеристики) '
        'одинаковы во всех документах — без противоречий и плейсхолдеров [НЕТ НА ЭТИКЕТКЕ].',
        f'07_designer_tz_{art}.pptx содержит по одному слайду на каждую карточку '
        'с вставленными фото конкурентов (9+ слайдов).',
        f'ТЗ_Фотограф_{art}.docx оформлен по Шаблон_ТЗ.docx: '
        'разделы 1–5, таблицы вводных данных и этапов заполнены, блоки съёмки именованы.',
        f'08_listing_{art}.docx содержит 3 варианта названия (до 200 символов), '
        'описание 1 500–3 000 символов с вписанными SEO-запросами, '
        'полные характеристики и обоснованную рекомендуемую цену.',
        'Рекомендованная цена обоснована: break-even посчитан, маржа при рекомендуемой цене > 0.',
        'Все поля [УТОЧНИТЬ У ПРОДАВЦА] переданы списком оператору.',
    ]
    for c in conditions:
        bullet(doc, c)

    # ── Доп: структура папки ───────────────────────────────────────────────
    doc.add_paragraph()
    heading(doc, 'Структура папки продукта', level=2)
    tree = (
        f'{base_folder}\n'
        f'  ├── 01_card_{art}.docx\n'
        f'  ├── 01a_tz_{art}.docx\n'
        f'  ├── 02_sales_history_{art}.docx\n'
        f'  ├── 03_competitors_{art}.xlsx\n'
        f'  ├── 03b_pricing_strategy_{art}.docx\n'
        f'  ├── 04_pains_{art}.docx\n'
        f'  ├── 05_seo_{art}.docx / .xlsx\n'
        f'  ├── 06_visual_{art}.docx\n'
        f'  ├── 07_designer_tz_{art}.docx / .pptx\n'
        f'  ├── ТЗ_Фотограф_{art}.docx\n'
        f'  ├── 08_listing_{art}.docx\n'
        f'  ├── Новая этикетка/\n'
        f'  ├── Конкурирующие товары/\n'
        f'  ├── Фото карточек конкурентов/\n'
        f'  └── Поисковые запросы/'
    )
    p_tree = doc.add_paragraph()
    p_tree.paragraph_format.space_after = Pt(4)
    p_tree.add_run(tree).font.size = Pt(9)

    # ── Доп: характеристики товара (если переданы) ─────────────────────────
    if extra_chars:
        doc.add_paragraph()
        heading(doc, 'Ключевые характеристики товара (из этикетки)', level=2)
        for k, v in extra_chars.items():
            kv(doc, k, str(v), size=9)

    doc.save(str(out_path))
    print(f'Сохранено: {out_path}  ({out_path.stat().st_size:,} байт)')
    return out_path


# ─── CLI ────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Генератор ТЗ по Шаблон_ТЗ.docx')
    parser.add_argument('--art',          required=True,  help='Артикул товара, напр. 770758')
    parser.add_argument('--brand',        required=True,  help='Бренд, напр. ifoamHOME')
    parser.add_argument('--name',         required=True,  help='Slug папки, напр. КондиционерBalmy_РайскиеЦветы')
    parser.add_argument('--product_name', required=False, default=None,
                        help='Читаемое название, напр. "balmy Кондиционер «Райские цветы»"')
    parser.add_argument('--volume',       required=False, default='',
                        help='Объём, напр. "1 л"')
    parser.add_argument('--date',         required=False, default=None,
                        help='Дата ДД.ММ.ГГГГ; если не указана — сегодня')
    args = parser.parse_args()

    from datetime import date
    date_str = args.date or date.today().strftime('%d.%m.%Y')
    product_name = args.product_name or f'{args.brand} {args.name}'

    folder = Path(f'C:/Users/1/Downloads/{args.brand}_{args.name}_{args.art}')

    build_product_tz(
        folder=folder,
        art=args.art,
        brand=args.brand,
        name=args.name,
        product_name=product_name,
        volume=args.volume,
        date_str=date_str,
    )
